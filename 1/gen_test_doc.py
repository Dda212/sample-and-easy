# -*- coding: utf-8 -*-
"""生成医院门诊预约管理系统测试文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)

# ===== 封面标题 =====
title = doc.add_heading('医院门诊预约管理系统', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('测 试 报 告')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(18)
sub.runs[0].font.bold = True
sub.runs[0].font.name = '黑体'
sub.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('项目名称：医院门诊预约管理系统\n').font.size = Pt(12)
info.add_run('文档版本：V1.0\n').font.size = Pt(12)
info.add_run('测试日期：2026-09-01\n').font.size = Pt(12)
info.add_run('测试环境：Windows 11 / Qt 6.11.1 MinGW 64-bit / C++17').font.size = Pt(12)

doc.add_page_break()

# ===== 1. 测试目的 =====
doc.add_heading('1. 测试目的', level=1)
doc.add_paragraph(
    '本文档对"医院门诊预约管理系统"进行全面功能测试，验证系统各模块是否满足设计需求，'
    '发现并记录潜在缺陷，确保系统在正常使用和边界场景下均能正确运行。测试范围覆盖程序启动、'
    '门诊管理、预约管理、查询功能、数据持久化、界面交互及输入校验七大模块。'
)

# ===== 2. 测试环境 =====
doc.add_heading('2. 测试环境', level=1)
env_table = doc.add_table(rows=6, cols=2, style='Table Grid')
env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
env_data = [
    ('操作系统', 'Windows 11'),
    ('开发框架', 'Qt 6.11.1 (MinGW 64-bit)'),
    ('编译器', 'GCC 13.1.0 (MinGW-w64)'),
    ('C++ 标准', 'C++17'),
    ('构建工具', 'qmake + mingw32-make'),
    ('测试数据', 'hospital_data.txt（2个门诊，2条预约记录）'),
]
for i, (k, v) in enumerate(env_data):
    env_table.rows[i].cells[0].text = k
    env_table.rows[i].cells[1].text = v
    env_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

# ===== 3. 测试范围 =====
doc.add_heading('3. 测试范围与模块划分', level=1)
scope_table = doc.add_table(rows=8, cols=3, style='Table Grid')
scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = scope_table.rows[0].cells
hdr[0].text = '模块编号'; hdr[1].text = '模块名称'; hdr[2].text = '测试内容'
for c in hdr:
    c.paragraphs[0].runs[0].font.bold = True
scope_data = [
    ('M1', '程序启动与初始化', '启动无崩溃、查询区默认隐藏、门诊数据自动加载、占位文字显示'),
    ('M2', '门诊管理', '查看门诊列表、添加门诊（正常/重复号/必填项/电话格式）'),
    ('M3', '预约管理', '正常预约、无门诊提示、全满提示、重复预约检测、电话校验、必填项'),
    ('M4', '查询功能', '查询区显隐切换、按门诊查询、按电话查询（存在/不存在）、占位项清空'),
    ('M5', '数据持久化', '添加门诊后保存、预约后保存、重启数据不丢失、数据路径固定'),
    ('M6', '界面与交互', '退出确认、窗口缩放布局跟随、表格列宽自适应'),
    ('M7', '输入校验', '手机号格式校验、必填项非空校验、容量边界控制'),
]
for i, (a, b, c) in enumerate(scope_data, 1):
    scope_table.rows[i].cells[0].text = a
    scope_table.rows[i].cells[1].text = b
    scope_table.rows[i].cells[2].text = c

# ===== 4. 测试用例 =====
doc.add_heading('4. 测试用例与执行结果', level=1)
doc.add_paragraph('以下为全部测试用例，"实际结果"与"是否通过"在执行测试后回填。')

# 测试用例数据: (编号, 模块, 用例名称, 前置条件, 操作步骤, 预期结果)
test_cases = [
    # M1 启动与初始化
    ('TC001', 'M1', '程序正常启动', '工程已编译生成1.exe',
     '双击运行1.exe', '程序窗口正常显示，标题为"医院门诊预约管理系统"，无崩溃、无报错弹窗'),
    ('TC002', 'M1', '查询区域默认隐藏', '程序已启动',
     '观察主界面初始状态', '查询表格（姓名/电话/性别等列）、电话输入框、门诊下拉框均不可见'),
    ('TC003', 'M1', '门诊数据自动加载', 'hospital_data.txt中存在2条门诊记录',
     '启动程序后观察门诊表格', '门诊表格自动显示2条门诊数据（1523儿科、1522泌尿科），无需手动点击"查看所有门诊"'),
    ('TC004', 'M1', '电话输入框占位文字', '查询区域已展开',
     '观察电话输入框', '输入框内显示灰色提示文字"输入电话号码后按回车查询"'),

    # M2 门诊管理
    ('TC005', 'M2', '查看所有门诊', '程序已启动，存在门诊数据',
     '点击"查看所有门诊"按钮', '门诊表格刷新并显示全部门诊记录，列含门诊号/名称/医生工号/接诊时间/容量/已预约/剩余号源'),
    ('TC006', 'M2', '添加门诊-正常流程', '程序已启动',
     '点击"添加门诊"→填写门诊号(如9001)、门诊名称(如测试科)、容量20→点确定', '弹窗提示"门诊添加成功"，门诊表格新增一条记录，数据文件立即更新'),
    ('TC007', 'M2', '添加门诊-门诊号重复', '已存在门诊号1523',
     '添加门诊时门诊号填1523，其余正常填写→点确定', '弹窗提示"门诊号 1523 已存在！"，不添加新记录'),
    ('TC008', 'M2', '添加门诊-门诊号为空', '程序已启动',
     '添加门诊时门诊号留空，其余正常填写→点确定', '弹窗提示"门诊号不能为空！"，不关闭对话框'),
    ('TC009', 'M2', '添加门诊-门诊名称为空', '程序已启动',
     '添加门诊时门诊名称留空，其余正常填写→点确定', '弹窗提示"门诊名称不能为空！"，不关闭对话框'),
    ('TC010', 'M2', '添加门诊-联系电话格式错误', '程序已启动',
     '添加门诊时联系电话填"12345"（非11位手机号）→点确定', '弹窗提示"联系电话格式不正确，请输入有效的11位手机号码！"，不关闭对话框'),

    # M3 预约管理
    ('TC011', 'M3', '预约门诊-正常流程', '存在可预约门诊（剩余号源>0）',
     '点击"预约门诊"→选择门诊→填写姓名、11位手机号、性别、年龄、预约时间→点确定', '弹窗提示"预约成功"，该门诊已预约数+1、剩余号源-1，数据文件立即更新'),
    ('TC012', 'M3', '预约门诊-无门诊时提示', '门诊列表为空',
     '点击"预约门诊"按钮', '弹窗提示"当前没有门诊信息，请先添加门诊！"，不打开预约对话框'),
    ('TC013', 'M3', '预约门诊-全部门诊已满', '所有门诊剩余号源均为0',
     '点击"预约门诊"按钮', '弹窗提示"当前所有门诊均已满，无法预约！"，不打开预约对话框'),
    ('TC014', 'M3', '预约门诊-重复预约检测', '某门诊已有电话为13994852257的预约',
     '对同一门诊再次使用电话13994852257进行预约', '弹窗提示"该电话已预约过此门诊，不可重复预约！"，不添加预约记录'),
    ('TC015', 'M3', '预约门诊-电话格式错误', '存在可预约门诊',
     '预约时电话填"abc123"→点确定', '弹窗提示"请输入有效的11位手机号码！"，不关闭对话框'),
    ('TC016', 'M3', '预约门诊-必填项为空', '存在可预约门诊',
     '预约时姓名留空→点确定；或电话留空→点确定；或预约时间留空→点确定', '分别提示"姓名不能为空！"/"电话不能为空！"/"预约时间不能为空！"，不关闭对话框'),

    # M4 查询功能
    ('TC017', 'M4', '查询区显隐切换', '程序已启动，查询区默认隐藏',
     '点击"查询预约"按钮→再次点击', '第一次点击：查询区显示（表格+输入框+下拉框可见）；第二次点击：查询区隐藏'),
    ('TC018', 'M4', '按门诊下拉查询', '查询区已显示，存在门诊及预约数据',
     '在门诊下拉框中选择"儿科"', '查询表格显示儿科门诊的所有预约记录（姓名/电话/性别/年龄/预约时间/所属门诊）'),
    ('TC019', 'M4', '按电话查询-存在记录', '查询区已显示，存在电话13994852257的预约',
     '在电话输入框输入13994852257→按回车', '查询表格显示该电话对应的所有预约记录'),
    ('TC020', 'M4', '按电话查询-不存在记录', '查询区已显示',
     '在电话输入框输入不存在的号码18800000000→按回车', '弹窗提示"未找到该电话的预约记录。"，查询表格为空'),
    ('TC021', 'M4', '切回占位项时清空表格', '已按门诊查询显示了结果',
     '在门诊下拉框中选择"-- 按门诊查询 --"', '查询表格清空（行数为0），不残留之前的查询结果'),

    # M5 数据持久化
    ('TC022', 'M5', '添加门诊后立即保存', '程序已启动',
     '添加一个新门诊→不退出程序，直接查看hospital_data.txt', 'hospital_data.txt中已包含新添加的门诊记录（无需退出程序）'),
    ('TC023', 'M5', '预约成功后立即保存', '存在可预约门诊',
     '成功预约一条记录→不退出程序，直接查看hospital_data.txt', 'hospital_data.txt中已包含新的预约记录'),
    ('TC024', 'M5', '重启后数据不丢失', '已添加门诊和预约记录并保存',
     '关闭程序→重新启动', '门诊表格和预约数据与关闭前一致，无丢失'),
    ('TC025', 'M5', '数据文件路径固定', '程序已启动',
     '从不同工作目录启动程序（命令行/双击exe）', '数据文件始终保存在exe所在目录，不因启动目录不同而变化'),

    # M6 界面与交互
    ('TC026', 'M6', '退出系统-确认对话框', '程序已启动',
     '点击"退出"按钮', '弹出确认对话框"确定要退出系统吗？"；点"否"不退出，点"是"保存数据并关闭'),
    ('TC027', 'M6', '窗口缩放布局跟随', '程序已启动',
     '拖动窗口边缘放大/缩小', '按钮、表格、输入框等控件随窗口大小自适应调整，表格列宽自动拉伸，不出现内容遮挡'),

    # M7 输入校验（边界）
    ('TC028', 'M7', '容量边界-约满后无法预约', '某门诊容量为1，已有1条预约（剩余0）',
     '对该门诊再次发起预约', '该门诊不出现在预约下拉框中（已过滤已满门诊），无法选择预约'),
]

# 创建测试用例表格
tc_table = doc.add_table(rows=1 + len(test_cases), cols=8, style='Table Grid')
tc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['用例编号', '模块', '用例名称', '前置条件', '操作步骤', '预期结果', '实际结果', '是否通过']
for i, h in enumerate(headers):
    cell = tc_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)

for row_idx, (tc_id, mod, name, pre, steps, expected) in enumerate(test_cases, 1):
    row = tc_table.rows[row_idx].cells
    row[0].text = tc_id
    row[1].text = mod
    row[2].text = name
    row[3].text = pre
    row[4].text = steps
    row[5].text = expected
    row[6].text = ''  # 实际结果待填
    row[7].text = ''  # 是否通过待填
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# 设置列宽
col_widths = [Cm(1.5), Cm(1.0), Cm(2.5), Cm(2.5), Cm(3.5), Cm(3.5), Cm(2.5), Cm(1.2)]
for row in tc_table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

doc.add_page_break()

# ===== 5. 缺陷记录 =====
doc.add_heading('5. 缺陷记录', level=1)
doc.add_paragraph('测试过程中发现的缺陷记录如下（如无缺陷则填写"无"）。')
bug_table = doc.add_table(rows=1, cols=6, style='Table Grid')
bug_table.alignment = WD_TABLE_ALIGNMENT.CENTER
bug_headers = ['缺陷编号', '关联用例', '缺陷描述', '严重程度', '状态', '修复说明']
for i, h in enumerate(bug_headers):
    bug_table.rows[0].cells[i].text = h
    bug_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# ===== 6. 测试总结 =====
doc.add_heading('6. 测试总结', level=1)
doc.add_paragraph('6.1 测试执行概况')
summary_table = doc.add_table(rows=5, cols=2, style='Table Grid')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_data = [
    ('测试用例总数', str(len(test_cases))),
    ('通过用例数', ''),
    ('失败用例数', ''),
    ('发现缺陷数', ''),
    ('测试结论', ''),
]
for i, (k, v) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = k
    summary_table.rows[i].cells[1].text = v
    summary_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
doc.add_paragraph('6.2 测试结论与建议')
doc.add_paragraph('（测试完成后填写：系统整体质量评价、遗留问题、上线建议等）')

# 保存
output_path = r'C:\Users\17476\OneDrive\Desktop\demo\project-new\1\医院门诊预约管理系统_测试报告.docx'
doc.save(output_path)
print(f'测试文档已生成: {output_path}')
print(f'测试用例总数: {len(test_cases)}')
