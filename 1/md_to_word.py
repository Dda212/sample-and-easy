#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将对话记录 Markdown 转换为格式规范的 Word 文档"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_PATH = r"C:\Users\17476\OneDrive\Desktop\demo\project-new\1\对话记录_项目全过程.md"
OUTPUT_PATH = r"C:\Users\17476\OneDrive\Desktop\demo\project-new\1\对话记录_项目全过程.docx"

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

for level, size in [(1, 16), (2, 14), (3, 12)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.paragraph_format.space_before = Pt(14 if level == 1 else 11 if level == 2 else 9)
    hs.paragraph_format.space_after = Pt(6 if level == 1 else 5 if level == 2 else 4)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge in ['top', 'bottom', 'left', 'right']:
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_runs_with_format(paragraph, text):
    """处理行内格式：**粗体** 和 `代码`"""
    # 分割 **粗体** 和 `代码`
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 128)
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading(text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph_text(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    add_runs_with_format(p, text)
    return p


def add_list_item(text, ordered=False, number=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    prefix = f"{number}. " if ordered else "• "
    run = p.add_run(prefix)
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    add_runs_with_format(p, text)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    return p


def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 灰色背景
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shd)
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run('\n')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 30, 30)


def add_table_from_md(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs_with_format(p, h.strip())
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, 'D9D9D9')
        set_cell_border(cell)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_runs_with_format(p, val.strip())
            for run in p.runs:
                run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
            set_cell_border(cell)

    doc.add_paragraph()  # 表后空行


# ========== 解析 Markdown ==========
with open(INPUT_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_lines = []
ordered_counter = 0

while i < len(lines):
    line = lines[i].rstrip('\n')

    # 代码块
    if line.strip().startswith('```'):
        if in_code_block:
            add_code_block(code_lines)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # 空行
    if not line.strip():
        i += 1
        continue

    # 水平线
    if line.strip() == '---':
        # 用一个空段落分隔
        doc.add_paragraph()
        i += 1
        continue

    # 标题
    if line.startswith('# '):
        add_heading(line[2:].strip(), 1)
        i += 1
        continue
    if line.startswith('## '):
        add_heading(line[3:].strip(), 2)
        i += 1
        continue
    if line.startswith('### '):
        add_heading(line[4:].strip(), 3)
        i += 1
        continue
    if line.startswith('#### '):
        add_heading(line[5:].strip(), 4)
        i += 1
        continue

    # 引用块
    if line.startswith('> '):
        add_quote(line[2:].strip())
        i += 1
        continue

    # 表格
    if '|' in line and line.strip().startswith('|'):
        # 收集表格行
        table_lines = []
        while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        # 解析表头
        headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
        # 跳过分隔行（第二行）
        data_lines = table_lines[2:] if len(table_lines) > 2 and re.match(r'^[\|\-:\s]+$', table_lines[1]) else table_lines[1:]
        rows = []
        for tl in data_lines:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            rows.append(cells)
        add_table_from_md(headers, rows)
        continue

    # 有序列表
    ordered_match = re.match(r'^(\d+)\.\s+(.*)', line)
    if ordered_match:
        num = int(ordered_match.group(1))
        text = ordered_match.group(2)
        add_list_item(text, ordered=True, number=num)
        i += 1
        continue

    # 无序列表
    if line.strip().startswith('- '):
        text = line.strip()[2:]
        add_list_item(text, ordered=False)
        i += 1
        continue

    # 普通段落
    add_paragraph_text(line.strip())
    i += 1

# 保存
doc.save(OUTPUT_PATH)
import os
print(f"文档已生成: {OUTPUT_PATH}")
print(f"文件大小: {os.path.getsize(OUTPUT_PATH)} bytes")
